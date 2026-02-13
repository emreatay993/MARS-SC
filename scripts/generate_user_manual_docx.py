"""
Generate MARS-SC user manual as a Word document with image placeholders.

This follows the same workflow style used in the sibling MARS_ project.
"""

from pathlib import Path


OUTPUT_PATH = Path(__file__).resolve().parents[1] / "MARS_SC_USER_MANUAL.docx"


def _require_docx():
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as exc:
        raise SystemExit(
            "python-docx is required. Install with: pip install python-docx"
        ) from exc
    return (
        Document,
        Inches,
        Pt,
        RGBColor,
        WD_ALIGN_PARAGRAPH,
        WD_TABLE_ALIGNMENT,
        qn,
        OxmlElement,
    )


def create_manual():
    (
        Document,
        Inches,
        Pt,
        RGBColor,
        WD_ALIGN_PARAGRAPH,
        WD_TABLE_ALIGNMENT,
        qn,
        OxmlElement,
    ) = _require_docx()

    def set_cell_shading(cell, color):
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), color)
        cell._tc.get_or_add_tcPr().append(shading)

    def add_placeholder_image(doc, caption, width_inches=5.6):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_shading(cell, "E8E8E8")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"\n\n[IMAGE PLACEHOLDER]\n\n{caption}\n\n")
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)
        cell.width = Inches(width_inches)
        doc.add_paragraph()

    def add_table(doc, headers, rows):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            set_cell_shading(header_cells[i], "5B9BD5")
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(10)
        for row_data in rows:
            row = table.add_row()
            for i, cell_text in enumerate(row_data):
                row.cells[i].text = str(cell_text)
        doc.add_paragraph()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_heading("MARS-SC: Solution Combination", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("User Manual")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(16)
        run.font.italic = True
    doc.add_paragraph()
    doc.add_paragraph()

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.add_run("Audience: ").bold = True
    intro.add_run("Structural/mechanical engineers using MARS-SC GUI\n\n")
    intro.add_run("Format: ").bold = True
    intro.add_run("Word-friendly chapters with image placeholders")

    doc.add_page_break()

    doc.add_heading("Table of Contents", 1)
    toc_items = [
        "Part I - Overview",
        "Part II - Named Selection Workflow",
        "Part III - Tooltip System",
        "Part IV - Quick Troubleshooting",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    doc.add_heading("Part I - Overview", 1)
    doc.add_heading("Chapter 1 - What Is MARS-SC", 2)
    doc.add_paragraph(
        "MARS-SC combines two static RST analyses using linear coefficients and "
        "provides envelope and per-combination visualization/export workflows."
    )
    add_table(
        doc,
        ["Capability", "Description"],
        [
            ["Two-analysis combination", "Combine Analysis 1 and Analysis 2 load-step results"],
            ["Named-selection scoping", "Scope solves to selected named selection nodes"],
            ["Display workflows", "Envelope/single-combination visualization and export"],
            ["Guided UI", "Detailed tooltips and global tooltip on/off control"],
        ],
    )
    add_placeholder_image(doc, "Main window with Solver and Display tabs")

    doc.add_page_break()

    doc.add_heading("Part II - Named Selection Workflow", 1)
    doc.add_heading("Chapter 2 - Named Selection Source Modes", 2)
    add_table(
        doc,
        ["Mode", "What is listed"],
        [
            ["Common (A1 & A2)", "Intersection of names in both analyses"],
            ["Analysis 1 (Base)", "All names from Analysis 1"],
            ["Analysis 2 (Combine)", "All names from Analysis 2"],
        ],
    )
    doc.add_paragraph(
        "If the same named selection exists in both analyses with different nodes, "
        "MARS-SC uses Analysis 1 node content (base precedence)."
    )
    add_placeholder_image(doc, "Named Selection Source dropdown and list behavior")

    doc.add_page_break()

    doc.add_heading("Part III - Tooltip System", 1)
    doc.add_heading("Chapter 3 - Coverage and Global Toggle", 2)
    doc.add_paragraph(
        "Tooltips are available across Solver and Display tabs and can be toggled "
        "globally via View -> Enable Tooltips."
    )
    add_table(
        doc,
        ["Area", "Highlights"],
        [
            ["Solver tab", "Input buttons, named selection controls, import CSV, output options"],
            ["Display tab", "Load visualization file format guidance, view/export controls"],
            ["Global behavior", "Persistent on/off state using QSettings"],
            ["Styling", "Tooltip style aligned with sibling MARS_ palette"],
        ],
    )
    add_placeholder_image(doc, "Tooltip examples and View menu toggle")

    doc.add_page_break()

    doc.add_heading("Part IV - Quick Troubleshooting", 1)
    doc.add_heading("Chapter 4 - Common Questions", 2)
    doc.add_paragraph("Q: Why is named-selection list empty?")
    doc.add_paragraph(
        "A: Ensure both RST files are loaded and source mode has available names."
    )
    doc.add_paragraph("Q: Why tooltips do not appear?")
    doc.add_paragraph("A: Verify View -> Enable Tooltips is checked.")

    doc.save(str(OUTPUT_PATH))
    print(f"Manual generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    create_manual()
